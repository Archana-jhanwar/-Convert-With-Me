from flask import Flask, render_template, request, jsonify, send_from_directory
import os
import uuid
from PIL import Image
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from pypdf import PdfReader, PdfWriter
import img2pdf

app = Flask(__name__)

BASE_DIR = os.path.abspath(os.path.dirname(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
CONVERTED_FOLDER = os.path.join(BASE_DIR, 'converted')
TEMP_FOLDER = os.path.join(BASE_DIR, 'temp')

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['CONVERTED_FOLDER'] = CONVERTED_FOLDER

for folder in [UPLOAD_FOLDER, CONVERTED_FOLDER, TEMP_FOLDER]:
    os.makedirs(folder, exist_ok=True)

def cleanup_file(filepath):
    try:
        if filepath and os.path.exists(filepath):
            os.remove(filepath)
    except Exception:
        pass

def get_file_extension(filename):
    if '.' in filename:
        return filename.rsplit('.', 1)[1].lower()
    return ''

def convert_image_format(input_path, output_path, output_format):
    try:
        with Image.open(input_path) as img:

            format_map = {
                "jpg": "JPEG",
                "jpeg": "JPEG",
                "png": "PNG",
                "webp": "WEBP",
                "bmp": "BMP"
            }

            save_format = format_map.get(output_format.lower())

            if not save_format:
                return False, "Unsupported format"

            if save_format == "JPEG" and img.mode in ("RGBA", "P"):
                img = img.convert("RGB")

            img.save(output_path, save_format)

        return True, None

    except Exception as e:
        return False, str(e)

def images_to_pdf(input_paths, output_path):
    try:
        with open(output_path, "wb") as f:
            f.write(img2pdf.convert(input_paths))
        return True, None
    except Exception as e:
        return False, str(e)

def docx_to_pdf(input_path, output_path):
    try:
        doc = Document(input_path)
        doc_template = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        for paragraph in doc.paragraphs:
            text = paragraph.text
            if text.strip():
                story.append(Paragraph(text, styles['Normal']))
        doc_template.build(story)
        return True, None
    except Exception as e:
        return False, str(e)

def pdf_to_docx(input_path, output_path):
    try:
        reader = PdfReader(input_path)
        doc = Document()
        for page in reader.pages:
            text = page.extract_text()
            if text.strip():
                doc.add_paragraph(text)
        doc.save(output_path)
        return True, None
    except Exception as e:
        return False, str(e)
def merge_pdfs(input_paths, output_path):
    try:
        writer = PdfWriter()

        for pdf_path in input_paths:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                writer.add_page(page)

        with open(output_path, "wb") as f:
            writer.write(f)

        return True, None
    except Exception as e:
        return False, str(e)

def split_pdf(input_path, output_path_base):
    try:
        reader = PdfReader(input_path)
        output_paths = []
        for i, page in enumerate(reader.pages):
            writer = PdfWriter()
            writer.add_page(page)
            output_path = f"{output_path_base}_page_{i+1}.pdf"
            with open(output_path, 'wb') as f:
                writer.write(f)
            output_paths.append(output_path)
        return True, output_paths, None
    except Exception as e:
        return False, [], str(e)

def txt_to_pdf(input_path, output_path):
    try:
        with open(input_path, 'r', encoding='utf-8') as file:
            text = file.read()
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        for line in text.split('\n'):
            if line.strip():
                story.append(Paragraph(line, styles['Normal']))
        doc.build(story)
        return True, None
    except Exception as e:
        return False, str(e)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        if 'files[]' not in request.files:
            return jsonify({'success': False, 'message': 'No files uploaded'})
        
        files = request.files.getlist('files[]')
        
        if not files or all(file.filename == '' for file in files):
            return jsonify({'success': False, 'message': 'No files selected'})
        
        conversion_type = request.form.get('conversionType', '')
        
        if conversion_type in ['jpg', 'png', 'webp', 'bmp']:
            return convert_single_image(files, conversion_type)
        elif conversion_type == 'image_to_pdf':
            return convert_images_to_pdf(files)
        elif conversion_type == 'docx_to_pdf':
            return convert_docx_to_pdf(files)
        elif conversion_type == 'pdf_to_docx':
            return convert_pdf_to_docx(files)
        elif conversion_type == 'merge_pdfs':
            return merge_multiple_pdfs(files)
        elif conversion_type == 'split_pdf':
            return split_single_pdf(files)
        elif conversion_type == 'txt_to_pdf':
            return convert_txt_to_pdf(files)
        else:
            return jsonify({'success': False, 'message': 'Invalid conversion type'})
            
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)})

def save_uploaded_files(files):
    saved_paths = []
    for file in files:
        if file and file.filename:
            ext = get_file_extension(file.filename)
            filename = f"{uuid.uuid4()}.{ext}" if ext else f"{uuid.uuid4()}"
            filepath = os.path.join(TEMP_FOLDER, filename)
            file.save(filepath)
            saved_paths.append(filepath)
    return saved_paths

def convert_single_image(files, output_format):
    if len(files) != 1:
        return jsonify({'success': False, 'message': 'Please upload exactly one image'})
    
    file = files[0]
    ext = get_file_extension(file.filename)
    input_filename = f"{uuid.uuid4()}.{ext}" if ext else f"{uuid.uuid4()}"
    input_path = os.path.join(TEMP_FOLDER, input_filename)
    file.save(input_path)
    
    output_filename = f"converted_{uuid.uuid4()}.{output_format}"
    output_path = os.path.join(CONVERTED_FOLDER, output_filename)
    
    success, error = convert_image_format(input_path, output_path, output_format)
    cleanup_file(input_path)
    
    if success:
        return jsonify({
            'success': True,
            'message': 'Image converted successfully!',
            'downloadUrl': f"/download/{output_filename}",
            'filename': output_filename
        })
    else:
        return jsonify({'success': False, 'message': f'Conversion failed: {error}'})

def convert_images_to_pdf(files):
    if len(files) < 1:
        return jsonify({'success': False, 'message': 'Please upload at least one image'})
    
    input_paths = save_uploaded_files(files)
    output_filename = f"converted_{uuid.uuid4()}.pdf"
    output_path = os.path.join(CONVERTED_FOLDER, output_filename)
    
    success, error = images_to_pdf(input_paths, output_path)
    
    for path in input_paths:
        cleanup_file(path)
    
    if success:
        return jsonify({
            'success': True,
            'message': 'PDF created successfully!',
            'downloadUrl': f"/download/{output_filename}",
            'filename': output_filename
        })
    else:
        return jsonify({'success': False, 'message': f'Conversion failed: {error}'})

def convert_docx_to_pdf(files):
    if len(files) != 1 or not files[0].filename.endswith('.docx'):
        return jsonify({'success': False, 'message': 'Please upload exactly one DOCX file'})
    
    file = files[0]
    input_path = os.path.join(TEMP_FOLDER, f"{uuid.uuid4()}_{file.filename}")
    file.save(input_path)
    
    output_filename = f"converted_{uuid.uuid4()}.pdf"
    output_path = os.path.join(CONVERTED_FOLDER, output_filename)
    
    success, error = docx_to_pdf(input_path, output_path)
    cleanup_file(input_path)
    
    if success:
        return jsonify({
            'success': True,
            'message': 'DOCX converted to PDF successfully!',
            'downloadUrl': f"/download/{output_filename}",
            'filename': output_filename
        })
    else:
        return jsonify({'success': False, 'message': f'Conversion failed: {error}'})

def convert_pdf_to_docx(files):
    if len(files) != 1 or not files[0].filename.endswith('.pdf'):
        return jsonify({'success': False, 'message': 'Please upload exactly one PDF file'})
    
    file = files[0]
    input_path = os.path.join(TEMP_FOLDER, f"{uuid.uuid4()}_{file.filename}")
    file.save(input_path)
    
    output_filename = f"converted_{uuid.uuid4()}.docx"
    output_path = os.path.join(CONVERTED_FOLDER, output_filename)
    
    success, error = pdf_to_docx(input_path, output_path)
    cleanup_file(input_path)
    
    if success:
        return jsonify({
            'success': True,
            'message': 'PDF converted to DOCX successfully!',
            'downloadUrl': f"/download/{output_filename}",
            'filename': output_filename
        })
    else:
        return jsonify({'success': False, 'message': f'Conversion failed: {error}'})

def merge_multiple_pdfs(files):
    if len(files) < 2:
        return jsonify({'success': False, 'message': 'Please upload at least 2 PDF files'})
    
    input_paths = save_uploaded_files(files)
    output_filename = f"merged_{uuid.uuid4()}.pdf"
    output_path = os.path.join(CONVERTED_FOLDER, output_filename)
    
    success, error = merge_pdfs(input_paths, output_path)
    
    for path in input_paths:
        cleanup_file(path)
    
    if success:
        return jsonify({
            'success': True,
            'message': 'PDFs merged successfully!',
            'downloadUrl': f"/download/{output_filename}",
            'filename': output_filename
        })
    else:
        return jsonify({'success': False, 'message': f'Merge failed: {error}'})

def split_single_pdf(files):
    if len(files) != 1 or not files[0].filename.endswith('.pdf'):
        return jsonify({'success': False, 'message': 'Please upload exactly one PDF file'})
    
    file = files[0]
    input_path = os.path.join(TEMP_FOLDER, f"{uuid.uuid4()}_{file.filename}")
    file.save(input_path)
    
    base_output = os.path.join(CONVERTED_FOLDER, f"split_{uuid.uuid4()}")
    
    success, output_paths, error = split_pdf(input_path, base_output)
    cleanup_file(input_path)
    
    if success and output_paths:
        filenames = [os.path.basename(p) for p in output_paths]
        return jsonify({
            'success': True,
            'message': f'PDF split into {len(output_paths)} pages successfully!',
            'downloadUrls': [f"/download/{f}" for f in filenames],
            'filenames': filenames,
            'multiple': True
        })
    else:
        return jsonify({'success': False, 'message': f'Split failed: {error}'})

def convert_txt_to_pdf(files):
    if len(files) != 1 or not files[0].filename.endswith('.txt'):
        return jsonify({'success': False, 'message': 'Please upload exactly one TXT file'})
    
    file = files[0]
    input_path = os.path.join(TEMP_FOLDER, f"{uuid.uuid4()}_{file.filename}")
    file.save(input_path)
    
    output_filename = f"converted_{uuid.uuid4()}.pdf"
    output_path = os.path.join(CONVERTED_FOLDER, output_filename)
    
    success, error = txt_to_pdf(input_path, output_path)
    cleanup_file(input_path)
    
    if success:
        return jsonify({
            'success': True,
            'message': 'TXT converted to PDF successfully!',
            'downloadUrl': f"/download/{output_filename}",
            'filename': output_filename
        })
    else:
        return jsonify({'success': False, 'message': f'Conversion failed: {error}'})

@app.route('/download/<filename>')
def download_file(filename):

    filepath = os.path.join(CONVERTED_FOLDER, filename)

    if not os.path.exists(filepath):
        return "File not found", 404

    return send_from_directory(
        CONVERTED_FOLDER,
        filename,
        as_attachment=True
    )
if __name__ == '__main__':
    app.run()
    print("Starting Flask server...")

    app.run(debug=True, host="0.0.0.0", port=5000)

